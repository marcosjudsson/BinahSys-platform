import io
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def generate_word_document(markdown_text, title="Relatório BinahSys"):
    """
    Gera um arquivo Word (.docx) a partir de um texto Markdown simplificado.
    Retorna um buffer de memória (BytesIO).
    """
    doc = Document()
    
    # Título do Documento
    heading = doc.add_heading(title, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Normaliza quebras de linha
    lines = markdown_text.split('\n')

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # Títulos (Headings)
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        
        # Listas (Bullet Points)
        elif line.startswith('* ') or line.startswith('- '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
            _format_paragraph(p) # Aplica negrito se houver
            
        # Listas Numeradas (Simples detecção de "1. ")
        elif re.match(r'^\d+\.\s', line):
            text_content = re.sub(r'^\d+\.\s', '', line)
            p = doc.add_paragraph(text_content, style='List Number')
            _format_paragraph(p)

        # Parágrafo Normal
        else:
            p = doc.add_paragraph(line)
            _format_paragraph(p)

    # Salva em memória
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def _format_paragraph(paragraph):
    """
    Função auxiliar simples para processar negrito (**texto**) dentro de um parágrafo.
    Nota: Esta é uma implementação básica. O python-docx não suporta markdown nativo,
    então precisamos reconstruir o run para aplicar estilos.
    """
    # Se não tiver negrito, não faz nada (mantém performance)
    text = paragraph.text
    if '**' not in text:
        return

    # Limpa o parágrafo atual para reconstruir com formatação
    paragraph.clear()
    
    # Divide o texto por '**'
    # Ex: "Texto **negrito** final" -> ['Texto ', 'negrito', ' final']
    parts = text.split('**')
    
    for i, part in enumerate(parts):
        run = paragraph.add_run(part)
        # Se o índice for ímpar (1, 3, 5...), é a parte que estava entre ** **
        if i % 2 != 0:
            run.bold = True
