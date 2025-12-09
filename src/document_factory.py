import io
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from src.image_factory import render_pdf_page_to_image
from src.utils import debug_log

def _format_paragraph(paragraph):
    """
    Função auxiliar para aplicar negrito (**texto**) em parágrafos.
    """
    text = paragraph.text
    if '**' not in text:
        return

    paragraph.clear()
    parts = text.split('**')

    for i, part in enumerate(parts):
        run = paragraph.add_run(part)
        if i % 2 != 0: # Índices ímpares são o conteúdo entre asteriscos
            run.bold = True

def generate_word_document(markdown_text, title="Relatório BinahSys", citations=None):
    """
    Gera um DOCX a partir de markdown e anexa imagens das citações fornecidas.
    
    Args:
        markdown_text (str): O conteúdo do relatório.
        title (str): Título do documento.
        citations (list): Lista de dicts [{'uri': 'gs://...', 'title': '...'}].
    """
    doc = Document()

    # --- 1. Título e Cabeçalho ---
    heading = doc.add_heading(title, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- 2. Processamento do Texto Markdown ---
    lines = markdown_text.split('\n')

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # Headers
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        
        # Listas
        elif line.startswith('* ') or line.startswith('- '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
            _format_paragraph(p)
        elif re.match(r'^\d+\.\s', line):
            text_content = re.sub(r'^\d+\.\s', '', line)
            p = doc.add_paragraph(text_content, style='List Number')
            _format_paragraph(p)
        
        # Texto Normal
        else:
            p = doc.add_paragraph(line)
            _format_paragraph(p)

    # --- 3. Processamento de Evidências Visuais (Imagens) ---
    debug_log(f"DEBUG DOCX: Citações recebidas: {len(citations) if citations else 0}")
    if citations:
        debug_log(f"DEBUG DOCX: Conteúdo das citações: {citations}")
        # Filtra apenas citações do Google Storage (gs://)
        gs_citations = [c for c in citations if c.get('uri', '').startswith('gs://')]
        debug_log(f"DEBUG DOCX: Citações GS filtradas: {len(gs_citations)}")
        
        if gs_citations:
            doc.add_page_break()
            doc.add_heading('Anexo: Fontes & Evidências Visuais', level=1)
            doc.add_paragraph("Abaixo estão as capturas automáticas dos documentos citados neste relatório.")

            processed_uris = set()

            for cit in gs_citations:
                uri = cit.get('uri')
                title = cit.get('title', 'Documento Sem Título')
                
                # Evita duplicatas visuais
                if uri in processed_uris:
                    continue
                processed_uris.add(uri)

                # Cabeçalho da Imagem
                doc.add_heading(f"Fonte: {title}", level=3)
                doc.add_paragraph(f"URI: {uri}", style="Quote")

                try:
                    # CHAMADA À FÁBRICA DE IMAGENS
                    # Por padrão pegamos a página 1.
                    # Futuro: Se o RAG retornar page number, passar aqui.
                    debug_log(f"DEBUG DOCX: Tentando renderizar imagem para: {uri}")
                    image_bytes = render_pdf_page_to_image(uri, page_number=1)

                    if image_bytes:
                        debug_log(f"DEBUG DOCX: Imagem renderizada com sucesso. Tamanho: {len(image_bytes)} bytes")
                        # Converte bytes para stream compatível com python-docx
                        image_stream = io.BytesIO(image_bytes)
                        
                        try:
                            # Adiciona imagem centralizada com largura máxima de 6 polegadas
                            doc.add_picture(image_stream, width=Inches(6.0))
                            last_paragraph = doc.paragraphs[-1] 
                            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            
                            doc.add_paragraph("Visualização da Página 1 gerada automaticamente via BinahSys/Vertex AI.", style="Caption")
                            debug_log(f"DEBUG DOCX: Imagem inserida com sucesso no documento!")
                        except Exception as img_err:
                            debug_log(f"DEBUG DOCX: ERRO ao inserir imagem no DOCX: {img_err}", error=img_err)
                            doc.add_paragraph(f"[Erro ao inserir imagem: {str(img_err)}]", style="Intense Quote")
                    else:
                        debug_log(f"DEBUG DOCX: render_pdf_page_to_image retornou None para {uri}")
                        doc.add_paragraph("[Imagem indisponível ou erro na renderização]", style="Intense Quote")

                except Exception as e:
                    debug_log(f"Erro ao processar imagem para DOCX: {uri}", error=e)
                    doc.add_paragraph(f"[Erro de processamento visual: {str(e)}]")
                
                doc.add_paragraph("") # Espaçamento

    # --- 4. Finalização ---
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    debug_log(f"DEBUG DOCX: Documento finalizado. Tamanho do buffer: {buffer.getbuffer().nbytes} bytes")
    return buffer